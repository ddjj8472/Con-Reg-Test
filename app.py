import streamlit as st
import pandas as pd
import html
import io
from docx import Document

st.set_page_config(
    page_title="용인시 건축 조례 지원 플랫폼",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="expanded"
)

from widget_utils import inject_floating_button
from processor import handle_ai_analysis, generate_civil_document
from style import apply_custom_style
from components import render_user_message, render_ai_report
from storage import load_history, save_history

inject_floating_button()


def get_query_param(name):
    value = st.query_params.get(name, None)
    if isinstance(value, list):
        return value[0] if value else None
    return value


def clear_query_params_and_rerun():
    st.query_params.clear()
    st.rerun()


qp_user = get_query_param("_user")
if "user_id" not in st.session_state:
    st.session_state.user_id = qp_user or "guest"
elif st.session_state.user_id == "guest" and qp_user:
    st.session_state.user_id = qp_user

if "chat_history" not in st.session_state:
    st.session_state.chat_history = load_history(st.session_state.user_id)
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False
if "dark_mode_toggle" not in st.session_state:
    st.session_state.dark_mode_toggle = st.session_state.dark_mode
if "selected_index" not in st.session_state:
    st.session_state.selected_index = None
if "current_page" not in st.session_state:
    st.session_state.current_page = "main"
if "qna_list" not in st.session_state:
    st.session_state.qna_list = []


def sync_dark_mode():
    st.session_state.dark_mode = st.session_state.dark_mode_toggle


def persist_history():
    save_history(st.session_state.chat_history, st.session_state.user_id)


# HTML 대화 메뉴 액션 처리
try:
    open_chat_idx = get_query_param("open_chat")
    pin_chat_idx = get_query_param("pin_chat")
    delete_chat_idx = get_query_param("delete_chat")
    rename_chat_idx = get_query_param("rename_chat")
    new_chat_title = get_query_param("new_title")

    if open_chat_idx is not None:
        idx = int(open_chat_idx)
        if 0 <= idx < len(st.session_state.chat_history):
            st.session_state.selected_index = idx
            st.session_state.current_page = "main"
        clear_query_params_and_rerun()

    if pin_chat_idx is not None:
        idx = int(pin_chat_idx)
        if 0 <= idx < len(st.session_state.chat_history):
            st.session_state.chat_history[idx]["pinned"] = not st.session_state.chat_history[idx].get("pinned", False)
            persist_history()
        clear_query_params_and_rerun()

    if delete_chat_idx is not None:
        idx = int(delete_chat_idx)
        if 0 <= idx < len(st.session_state.chat_history):
            st.session_state.chat_history.pop(idx)
            if st.session_state.selected_index == idx:
                st.session_state.selected_index = None
            elif st.session_state.selected_index is not None and st.session_state.selected_index > idx:
                st.session_state.selected_index -= 1
            persist_history()
        clear_query_params_and_rerun()

    if rename_chat_idx is not None:
        idx = int(rename_chat_idx)
        if 0 <= idx < len(st.session_state.chat_history):
            title = (new_chat_title or "").strip()
            if title:
                st.session_state.chat_history[idx]["title"] = title
                persist_history()
        clear_query_params_and_rerun()
except Exception as e:
    st.error(f"대화 메뉴 처리 중 오류 발생: {e}")


apply_custom_style(st.session_state.dark_mode)

st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
html, body, .stApp, p, h1, h2, h3, h4, h5, h6, label, input, textarea, div, button {
    font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, sans-serif;
}
div[data-testid="stButton"] button, div[data-testid="stFormSubmitButton"] button {
    border-radius: 8px !important;
    font-weight: 600 !important;
}
.chat-history-card {
    position: relative;
    display: flex;
    align-items: center;
    justify-content: space-between;
    width: 100%;
    margin-bottom: 8px;
    border: 1px solid #DEE2E6;
    border-radius: 10px;
    background: #FFFFFF;
    box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    overflow: visible;
}
.chat-history-card:hover { background: #F8F9FA; border-color: #C9D4E2; }
.chat-open-form { flex: 1; margin: 0; min-width: 0; }
.chat-open-btn {
    width: 100%; border: none; background: transparent; text-align: left;
    padding: 10px 6px 10px 10px; font-size: 13px; font-weight: 650;
    color: #222; cursor: pointer; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
}
.chat-menu { position: relative; flex: 0 0 34px; padding-right: 5px; }
.chat-menu summary {
    list-style: none; cursor: pointer; width: 28px; height: 28px; border-radius: 7px;
    display: flex; align-items: center; justify-content: center;
    font-size: 20px; font-weight: 900; color: #555; user-select: none;
}
.chat-menu summary::-webkit-details-marker { display: none; }
.chat-menu summary:hover { background: #ECEFF3; }
.chat-menu-panel {
    position: absolute; right: 0; top: 32px; z-index: 999999; width: 196px;
    padding: 10px; border: 1px solid #D7DEE8; border-radius: 10px;
    background: white; box-shadow: 0 8px 20px rgba(0,0,0,0.14);
}
.chat-menu-label { font-size: 12px; font-weight: 800; color: #555; margin-bottom: 6px; }
.chat-title-input {
    box-sizing: border-box; width: 100%; padding: 7px 8px; margin-bottom: 6px;
    border: 1px solid #CED4DA; border-radius: 7px; font-size: 12px;
}
.chat-menu-btn {
    width: 100%; padding: 7px 8px; margin: 3px 0; border: 1px solid #DEE2E6;
    border-radius: 7px; background: #F8F9FA; color: #222; font-size: 12px;
    font-weight: 700; cursor: pointer; text-align: left;
}
.chat-menu-btn:hover { background: #E9ECEF; }
.chat-menu-btn.danger { color: #C92A2A; border-color: #FFC9C9; background: #FFF5F5; }
</style>
""", unsafe_allow_html=True)


dialog_decorator = st.dialog if hasattr(st, "dialog") else st.experimental_dialog

@dialog_decorator("🔍 대화기록 검색", width="large")
def open_history_search_dialog():
    search_query = st.text_input("검색어 입력", placeholder="예: 일조권, 건폐율...", key="dialog_history_search_input")
    query = search_query.strip().lower()
    if not st.session_state.chat_history:
        st.caption("저장된 대화기록이 없습니다.")
        return
    results = []
    for idx, chat in enumerate(st.session_state.chat_history):
        text = chat.get("title", "") + " "
        for msg in chat.get("messages", []):
            text += msg.get("query", "") + " " + msg.get("response", "") + " "
        if not query or query in text.lower():
            results.append((idx, chat))
    st.caption(f"검색 결과 {len(results)}개" if query else "최근 대화")
    if not results:
        st.warning("검색 결과가 없습니다.")
        return
    with st.container(height=400, border=False):
        for idx, chat in reversed(results[-20:]):
            title = chat.get("title", "새 대화")
            time_str = chat.get("updated_at", chat.get("created_at", ""))
            if st.button(f"💬 {title[:25]}...\n\n{time_str}", key=f"dialog_chat_{idx}", use_container_width=True):
                st.session_state.selected_index = idx
                st.session_state.current_page = "main"
                st.rerun()
            st.markdown("---")


with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/1048/1048978.png", width=50)
    st.title("플랫폼 제어")

    with st.expander("👤 실무자 시스템 인증", expanded=(st.session_state.user_id == "guest")):
        if st.session_state.user_id == "guest":
            tabs = st.tabs(["로그인", "회원가입"])
            with tabs[0]:
                login_id = st.text_input("아이디", key="login_id")
                login_pw = st.text_input("비밀번호", type="password", key="login_pw")
                if st.button("로그인", use_container_width=True):
                    from storage import authenticate_user
                    if authenticate_user(login_id.strip(), login_pw.strip()):
                        st.session_state.user_id = login_id.strip()
                        st.session_state.chat_history = load_history(st.session_state.user_id)
                        st.toast(f"환영합니다, {st.session_state.user_id}님!", icon="👋")
                        st.rerun()
                    else:
                        st.error("정보가 일치하지 않습니다.")
            with tabs[1]:
                reg_id = st.text_input("새 아이디", key="reg_id")
                reg_pw = st.text_input("비밀번호", type="password", key="reg_pw")
                if st.button("가입하기", use_container_width=True):
                    from storage import check_id_exists, register_user
                    if check_id_exists(reg_id.strip()):
                        st.error("이미 존재하는 아이디입니다.")
                    elif register_user(reg_id.strip(), reg_pw.strip()):
                        st.toast("가입 성공! 로그인해주세요.", icon="✅")
        else:
            st.success(f"현재 접속: **{st.session_state.user_id}**")
            if st.button("안전 로그아웃", use_container_width=True):
                st.session_state.user_id = "guest"
                st.session_state.chat_history = load_history("guest")
                st.session_state.selected_index = None
                st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("📌 메인 메뉴")
    if st.button("🏠 메인화면 (AI 챗봇)", use_container_width=True):
        st.session_state.current_page = "main"; st.rerun()
    if st.button("📝 민원 양식 자동생성", use_container_width=True):
        st.session_state.current_page = "doc_gen"; st.rerun()
    if st.button("💡 FAQ & Q&A 게시판", use_container_width=True):
        st.session_state.current_page = "qna"; st.rerun()
    if st.button("🗺️ 플랫폼 사이트맵", use_container_width=True):
        st.session_state.current_page = "sitemap"; st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    st.toggle("🌙 다크 모드", key="dark_mode_toggle", on_change=sync_dark_mode)
    st.markdown("---")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("➕ 새 대화", use_container_width=True):
            st.session_state.selected_index = None
            st.session_state.current_page = "main"
            st.rerun()
    with col2:
        if st.button("🔍 검색", use_container_width=True):
            open_history_search_dialog()

    st.subheader("📁 대화 이력")
    with st.container(height=250, border=True):
        if st.session_state.chat_history:
            items = list(enumerate(st.session_state.chat_history))
            items.sort(key=lambda item: (item[1].get("pinned", False), item[1].get("updated_at", item[1].get("created_at", ""))), reverse=True)
            user_safe = html.escape(st.session_state.user_id, quote=True)
            for actual_index, chat in items:
                time_str = chat.get("updated_at", chat.get("created_at", "00-00 00:00"))[5:16]
                title = chat.get("title", "새 대화")
                short_title = title[:14] + ".." if len(title) > 14 else title
                pin_mark = "📌 " if chat.get("pinned", False) else ""
                pin_label = "고정 해제" if chat.get("pinned", False) else "상단 고정"
                time_safe = html.escape(time_str, quote=True)
                title_safe = html.escape(title, quote=True)
                short_safe = html.escape(short_title, quote=True)
                pin_label_safe = html.escape(pin_label, quote=True)
                st.markdown(f"""<div class="chat-history-card"><form class="chat-open-form" method="get" action=""><input type="hidden" name="_user" value="{user_safe}"><input type="hidden" name="open_chat" value="{actual_index}"><button class="chat-open-btn" type="submit">{pin_mark}🕒 {time_safe} | {short_safe}</button></form><details class="chat-menu"><summary>⋯</summary><div class="chat-menu-panel"><div class="chat-menu-label">대화 옵션</div><form method="get" action=""><input type="hidden" name="_user" value="{user_safe}"><input type="hidden" name="rename_chat" value="{actual_index}"><input class="chat-title-input" type="text" name="new_title" value="{title_safe}"><button class="chat-menu-btn" type="submit">제목 저장</button></form><form method="get" action=""><input type="hidden" name="_user" value="{user_safe}"><input type="hidden" name="pin_chat" value="{actual_index}"><button class="chat-menu-btn" type="submit">{pin_label_safe}</button></form><form method="get" action="" onsubmit="return confirm('이 대화 기록을 삭제하시겠습니까?');"><input type="hidden" name="_user" value="{user_safe}"><input type="hidden" name="delete_chat" value="{actual_index}"><button class="chat-menu-btn danger" type="submit">삭제</button></form></div></details></div>""", unsafe_allow_html=True)
        else:
            st.caption("저장된 기록이 없습니다.")

    if st.session_state.chat_history:
        if st.button("🗑️ 전체 기록 삭제", type="primary", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.selected_index = None
            from storage import clear_history
            clear_history(st.session_state.user_id)
            st.toast("모든 기록이 삭제되었습니다.", icon="🗑️")
            st.rerun()


if st.session_state.current_page == "main":
    st.title("🏢 건축 조례 및 법령 해석 AI")
    st.markdown("<p style='color: #666; font-size: 1.1em;'>건축사, 시공사, 인허가 담당자의 신속한 의사결정을 돕는 심층 규제 분석 엔진입니다.</p>", unsafe_allow_html=True)
    chat_box = st.container(height=500, border=False)
    user_query = st.chat_input("예: 용인시 처인구 자연녹지지역의 건폐율과 용적률 기준은?")

    with chat_box:
        if st.session_state.selected_index is not None and 0 <= st.session_state.selected_index < len(st.session_state.chat_history):
            idx = st.session_state.selected_index
            selected_chat = st.session_state.chat_history[idx]
            st.info(f"📅 과거 대화 열람 중: **{selected_chat.get('title', '새 대화')}**")
            for msg in selected_chat.get("messages", []):
                render_user_message(msg.get("query", ""))
                render_ai_report(msg.get("response", ""))
            if st.button("닫기 및 새 질문하기", use_container_width=True):
                st.session_state.selected_index = None
                st.rerun()
        else:
            st.image("https://images.unsplash.com/photo-1486406146926-c627a92ad1ab?q=80&w=2070&auto=format&fit=crop", use_container_width=True)
            st.markdown("""
            <div style="text-align:center; padding: 30px; background: rgba(128,128,128,0.05); border-radius: 12px; margin-top: 15px;">
                <h3 style="color: #0b459c; margin-bottom: 8px;">어떤 규제를 검토해 드릴까요?</h3>
                <p style="color: #555; font-size: 14px; line-height: 1.6;">경기도/용인시 조례 및 상위 법령 데이터베이스를 기반으로 분석합니다.</p>
            </div>
            """, unsafe_allow_html=True)

        if user_query:
            render_user_message(user_query)
            with st.status("🔍 법률 시맨틱 엔진 가동 중...", expanded=True) as status:
                try:
                    st.write("🛰️ 조항 필터링 및 교차 검증 진행 중...")
                    response_text = handle_ai_analysis(user_query)
                    if st.session_state.chat_history:
                        st.session_state.selected_index = len(st.session_state.chat_history) - 1
                    status.update(label="✅ 심층 분석 완료", state="complete")
                    render_ai_report(response_text)
                    st.toast("분석이 완료되었습니다.", icon="✨")
                except Exception as e:
                    status.update(label="❌ 시스템 에러 발생", state="error")
                    st.error(f"오류가 발생했습니다: {str(e)}")
            st.rerun()

elif st.session_state.current_page == "doc_gen":
    st.title("📝 맞춤형 건축 민원 양식 자동완성")
    st.info("복잡한 민원 내용을 입력하면 AI가 용인시 행정 양식에 맞춰 문서를 작성합니다. 법령 검토는 챗봇을 이용해 주세요.")
    required_docs = {
        "건축허가 관련": ["건축허가 신청서", "배치도", "평면도", "토지이용계획확인서", "건축계획서"],
        "건축선 문의": ["대지 위치도", "토지이용계획확인서", "현장 사진"],
        "일조권 민원": ["현장 사진", "건축물 배치도", "피해 설명 자료"],
        "불법건축물 신고": ["현장 사진", "위치도", "불법사항 설명자료"],
        "용도변경 문의": ["건축물대장", "평면도", "용도변경 계획서"],
        "주차장 기준 문의": ["배치도", "주차계획도", "건축 개요"],
        "건축물 해석 문의": ["질의서", "관련 도면", "현장 사진"],
        "기타": ["신분증", "민원 설명자료"]
    }
    department_map = {"건축허가 관련": "건축허가과", "건축선 문의": "건축과", "일조권 민원": "건축과", "불법건축물 신고": "건축과", "용도변경 문의": "건축허가과", "주차장 기준 문의": "교통정책과", "건축물 해석 문의": "건축과", "기타": "민원여권과"}
    col1, col2 = st.columns(2)
    with col1:
        civil_type = st.selectbox("📌 민원 유형 선택", list(required_docs.keys()))
    with col2:
        site_address = st.text_input("📍 대상 건축물 주소", placeholder="예: 경기도 용인시 처인구 중부대로 1199")
    civil_content = st.text_area("✏️ 민원 상세 내용", height=150)
    if st.button("✨ AI 민원 양식 생성하기", use_container_width=True, type="primary"):
        if not site_address or not civil_content:
            st.error("주소와 민원 내용을 모두 입력해주세요.")
        else:
            result = generate_civil_document(civil_type, site_address, civil_content)
            st.subheader("📄 생성된 민원서")
            st.markdown(f"<div style='padding:20px; border:1px solid #ddd; border-radius:8px; background:rgba(0,0,0,0.02);'>{result}</div>", unsafe_allow_html=True)
            tabs = st.tabs(["📎 필요 서류", "🏢 담당 부서", "📥 파일 다운로드"])
            with tabs[0]:
                for doc_name in required_docs.get(civil_type, required_docs["기타"]):
                    st.write(f"✔️ {doc_name}")
            with tabs[1]:
                st.info(f"📌 용인시청 또는 관할 구청 **{department_map.get(civil_type, '민원여권과')}** 문의 요망")
            with tabs[2]:
                doc = Document()
                doc.add_heading('용인시 건축 행정 민원서', level=1)
                doc.add_paragraph(f"민원 유형: {civil_type}\n대상 주소: {site_address}")
                doc.add_heading('민원 내용', level=2)
                doc.add_paragraph(result)
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button("💾 DOCX 양식 다운로드", buffer, "용인시_건축민원서.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif st.session_state.current_page == "qna":
    st.title("💡 커뮤니티 Q&A")
    st.write("플랫폼 사용법이나 애매한 규제 해석에 대해 질문을 남겨주세요.")
    q_title = st.text_input("제목")
    q_content = st.text_area("내용")
    if st.button("질문 등록", use_container_width=True):
        if q_title and q_content:
            st.session_state.qna_list.append({"title": q_title, "content": q_content, "status": "대기중", "answer": ""})
            st.rerun()
    for q in st.session_state.qna_list:
        with st.expander(f"[{q['status']}] {q['title']}"):
            st.markdown(f"**Q.** {q['content']}")

elif st.session_state.current_page == "sitemap":
    st.title("🗺️ 시스템 아키텍처 및 취급 데이터")
    st.write("본 플랫폼은 AI 엔진과 법규 DB를 연동하여 동작합니다.")
    df_ord = pd.DataFrame([
        ["경기도", "경기도 건축 조례", "건축법"],
        ["경기도", "경기도 건축기본조례", "건축기본법"],
        ["용인시", "용인시 건축 조례", "건축법"],
        ["용인시", "용인시 도시계획 조례", "국토계획법"]
    ], columns=["지자체", "조례명", "근거법률"])
    st.dataframe(df_ord, hide_index=True, use_container_width=True)
